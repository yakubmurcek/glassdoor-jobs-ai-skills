import re
from docx import Document

def generate_docx():
    doc = Document()
    
    with open('docs/04_metodika_struktura_datasetu.md', 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            content = line[2:]
            
            # Simple bold search for **text**: text
            m = re.match(r'\*\*(.*?)\*\*(.*)', content)
            if m:
                b_text = m.group(1).replace('`', '')
                rest_text = m.group(2).replace('`', '')
                
                run = p.add_run(b_text)
                run.bold = True
                p.add_run(rest_text)
            else:
                p.add_run(content.replace('`', ''))
        else:
            doc.add_paragraph(line.replace('`', ''))

    output_path = 'docs/04_metodika_struktura_datasetu.docx'
    doc.save(output_path)
    print(f"Byl vytvořen soubor s nativními Word styly: {output_path}")

if __name__ == '__main__':
    generate_docx()
